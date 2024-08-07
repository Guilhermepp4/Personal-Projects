from flask import Flask, render_template, request, send_file
import os
import tempfile
import zipfile
from docx import Document
from win32com import client
import pythoncom
from concurrent.futures import ThreadPoolExecutor

app = Flask(__name__)
FILES_DIR = './files/'

def process_file(file_name, replaceText1, replaceText2, replaceText3, footer_text, village_Text):
    file_path = os.path.join(FILES_DIR, file_name)
    
    if not os.path.isfile(file_path):
        return None, f"File not found: {file_name}"

    try:
        # Load and process the document
        document = Document(file_path)

        start_phrase = 'execução relativo à “'
        end_phrase = '” e cujo Dono de Obra '
        start_phrase2 = "Dono de Obra é "
        end_phrase2 = ", observa as normas legais "
        start_phrase3 = "Braga, "
        end_phrase3 = " 202"

        for paragraph in document.paragraphs:
            text = paragraph.text

            # First replacement
            start_index = text.find(start_phrase)
            if start_index != -1:
                end_index = text.find(end_phrase, start_index + len(start_phrase))
                if end_index != -1:
                    end_index += len(end_phrase)
                    new_text = text[:start_index] + start_phrase + replaceText1 + end_phrase + text[end_index:]
                    paragraph.text = new_text
                else:
                    paragraph.text = text[:start_index] + replaceText1 + text[start_index + len(start_phrase):]

            # Update paragraph text after first replacement
            text = paragraph.text

            # Second replacement
            start_index = text.find(start_phrase2)
            if start_index != -1:
                end_index = text.find(end_phrase2, start_index + len(start_phrase2))
                if end_index != -1:
                    end_index += len(end_phrase2)
                    new_text = text[:start_index] + start_phrase2 + replaceText2 + end_phrase2 + text[end_index:]
                    paragraph.text = new_text
                else:
                    paragraph.text = text[:start_index] + replaceText2 + text[start_index + len(start_phrase2):]
            
            # Update paragraph text after second replacement
            text = paragraph.text

            # Third replacement
            start_index = text.find(start_phrase3)
            if start_index != -1:
                end_index = text.find(end_phrase3, start_index + len(start_phrase3))
                if end_index != -1:
                    end_index += len(end_phrase3)
                    new_text = text[:start_index] + start_phrase3 + replaceText3 + end_phrase3 + text[end_index:]
                    paragraph.text = new_text
                else:
                    paragraph.text = text[:start_index] + replaceText3 + text[start_index + len(start_phrase3):]
            
            # Fourth replacement
            if "AQUI" in text:
                paragraph.text = text.replace("AQUI", "")
        
        # Update footer in each section
        run_Footer = False
        run_Village = False

        for section in document.sections:
            footer = section.footer
            for table in footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if paragraph.text.strip():  # Check if the cell already contains text
                                if not run_Footer:
                                    paragraph.clear()  # Clear existing text
                                    paragraph.add_run(footer_text)
                                    run_Footer = True
                                elif not run_Village:
                                    paragraph.clear()
                                    paragraph.add_run(village_Text)
                                    run_Village = True
                            if run_Footer and run_Village:
                                break
                    if run_Footer and run_Village:
                        break
                if run_Footer and run_Village:
                    break

        # Save the document as a temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_docx:
            document.save(temp_docx.name)
            temp_docx_path = temp_docx.name

        # Define the final PDF file name based on the uploaded file's original name
        original_filename = os.path.splitext(file_name)[0]
        final_pdf_filename = f"{original_filename}.pdf"
        pdf_path = os.path.join(tempfile.gettempdir(), final_pdf_filename)

        try:
            pythoncom.CoInitialize()  # Initialize COM library for thread
            word = client.Dispatch("Word.Application")
            doc = word.Documents.Open(temp_docx_path)
            
            doc.SaveAs(pdf_path, FileFormat=17)
            doc.Close()
            word.Quit()
        except Exception as e:
            return None, f"Error converting file to PDF: {e}"
        finally:
            os.remove(temp_docx_path)

        return pdf_path, None

    except Exception as e:
        return None, f"Error processing file {file_name}: {e}"

@app.route('/')
def home():
    # List files in the directory
    files = [f for f in os.listdir(FILES_DIR) if os.path.isfile(os.path.join(FILES_DIR, f))]
    return render_template('index.html', files=files)

@app.route('/replacementText', methods=['POST'])
def replacementForm():
    # Get the list of selected filenames
    selected_files = request.form.getlist('selected')

    if len(selected_files) == 0:
        return "You have to select at least one file", 400

    # Access replacement texts
    replaceText1 = request.form.get('replaceText1')
    replaceText2 = request.form.get('replaceText2')
    replaceText3 = request.form.get('replaceText3')
    footer_text = request.form.get('footerOwner')
    village_Text = request.form.get('footerLocal')

    processed_files = []
    errors = []

    with ThreadPoolExecutor(max_workers=4) as executor:
        futures = {executor.submit(process_file, file_name, replaceText1, replaceText2, replaceText3, footer_text, village_Text): file_name for file_name in selected_files}
        for future in futures:
            result, error = future.result()
            if error:
                errors.append(error)
            elif result:
                processed_files.append(result)

    if processed_files:
        zip_path = os.path.join(tempfile.gettempdir(), "processed_files.zip")
        with zipfile.ZipFile(zip_path, 'w') as zipf:
            for pdf_path in processed_files:
                zipf.write(pdf_path, os.path.basename(pdf_path))

        # Cleanup individual PDF files
        for pdf_path in processed_files:
            os.remove(pdf_path)

        return send_file(zip_path, as_attachment=True, download_name="processed_files.zip", mimetype="application/zip")

    error_message = "No files were processed. " + " ".join(errors) if errors else "No files were processed."
    return error_message, 400

if __name__ == '__main__':
    app.run(debug=True)
